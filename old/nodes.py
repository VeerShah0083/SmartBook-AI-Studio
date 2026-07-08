"""
nodes.py — All Six Agent Node Implementations
==============================================
Every node is a plain Python function that receives the LangGraph state dict
and returns a partial-update dict.  Human-in-the-Loop nodes call interrupt()
which pauses graph execution and waits for Command(resume=<value>).

NODE MAP (matches blueprint doc numbering):
  Node 1  →  scope_architect        — Interactive questionnaire
  Node 3  →  structural_scaffold    — Automated TOC builder
  Node 4a →  writer_gate            — Micro-approval style selection
  Node 2  →  targeted_researcher    — Automated RAG utility
  Node 4b →  writer_draft           — Full chapter draft generator
  Node 5  →  review_editor          — Binary HITL gatekeeper
  Node 6  →  doc_exporter           — DOCX / PDF compiler

Design notes vs original blueprint
-----------------------------------
• Node 4 is split into writer_gate + writer_draft to keep the
  HITL intercept and the bulk generation as distinct graph nodes.
• Node 2 is a proper graph node (not an inline tool call) so it
  benefits from checkpointing and can be independently retried.
• Questionnaire generation uses temperature=0 for determinism —
  critical because interrupt() re-runs the node on every resume;
  without this, different questions would appear each re-run.
• approved_chapters uses LangGraph's operator.add reducer so each
  approval appends without overwriting previous chapters.
• review_editor makes two sequential interrupt() calls when the
  user chooses REVISE (one for the decision, one for the feedback).
"""

from __future__ import annotations

import datetime
import json
import os
import re
from typing import Any

from anthropic import Anthropic
from langgraph.types import interrupt

from state import TextbookSystemState

# ─── LLM client ──────────────────────────────────────────────────────────────
_client = Anthropic()
MODEL   = "claude-sonnet-4-6"

# ─── Generic helpers ─────────────────────────────────────────────────────────

def _llm(system: str, user: str, *, max_tokens: int = 2048,
          temperature: float = 0.7) -> str:
    """Thin wrapper: call Claude and return the text response."""
    resp = _client.messages.create(
        model=MODEL,
        max_tokens=max_tokens,
        temperature=temperature,
        system=system,
        messages=[{"role": "user", "content": user}],
    )
    return resp.content[0].text


def _llm_json(system: str, user: str, *, max_tokens: int = 2048,
               temperature: float = 0.0) -> Any:
    """
    Call Claude expecting a JSON response.
    temperature=0 is the default for JSON calls that must be deterministic
    on node re-runs after interrupt() resumes.
    """
    suffix = (
        "\n\nCRITICAL: Respond with VALID JSON only. "
        "No markdown code fences. No preamble. No trailing commentary."
    )
    raw = _llm(system + suffix, user, max_tokens=max_tokens,
               temperature=temperature)

    text = raw.strip()
    # Strip markdown fences if the model added them anyway
    if text.startswith("```"):
        lines = text.splitlines()
        text = "\n".join(
            lines[1:-1] if lines[-1].strip() == "```" else lines[1:]
        )

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # Last resort: find the outermost JSON object / array
        m = re.search(r"[\[{].*[\]}]", text, re.DOTALL)
        if m:
            return json.loads(m.group())
        raise ValueError(f"Could not parse JSON.\nRaw response:\n{raw[:600]}")


# ═════════════════════════════════════════════════════════════════════════════
# NODE 1 — SCOPE & CURRICULUM ARCHITECT  (Interactive HITL)
# ═════════════════════════════════════════════════════════════════════════════

_QGEN_SYSTEM = """
You are an expert curriculum architect.

Given a subject and grade level, generate a structured questionnaire that will
capture everything needed to produce a high-quality academic textbook.

Return a JSON object with EXACTLY this structure:
{
  "question_count": <int 5-15>,
  "reasoning": "<1-2 sentences explaining why this many questions>",
  "questions": [
    {
      "id": "q1",
      "category": "<pedagogy|standards|tone|depth|assessment|audience|format>",
      "text": "<clear, jargon-free question — one sentence>",
      "options": ["<Option A>", "<Option B>", "<Option C>", "<Option D>"],
      "implications": [
        "<plain-English consequence of choosing A>",
        "<consequence of B>",
        "<consequence of C>",
        "<consequence of D>"
      ]
    }
  ]
}

Scale rules:
• Kindergarten / Early Elementary  →  5 questions  (minimal constraints)
• Upper Elementary / Middle School →  8 questions
• High School                      → 10 questions
• Undergraduate / Adult Technical  → 12-15 questions

Cover these topic categories across your questions (not all categories
are mandatory for every level):
  pedagogy, standards, tone, depth, assessment, audience, format
"""

_SCOPE_SYSTEM = """
You are a senior curriculum architect writing a formal Project Scope Profile.

Based on the subject, grade level, and questionnaire answers, produce a
structured scope document.  Be specific and actionable — every downstream
agent reads this document to constrain content generation.

Include ALL of the following sections:
1. Proposed textbook title and subtitle
2. Target audience description (age, background assumptions, accessibility)
3. Core learning pillars  (5–7 high-level themes / big ideas)
4. Reading level targets  (Flesch-Kincaid grade estimate, vocabulary tier)
5. Curriculum standards alignment  (list each standard)
6. Pedagogical philosophy and primary teaching methods
7. Assessment philosophy  (formative vs summative balance)
8. Recommended chapter count  (justify the number)
9. Tone, voice, and style guidelines  (concrete examples)
10. Content inclusion / exclusion boundaries  (explicit rules)
"""


def scope_architect(state: TextbookSystemState) -> dict:
    """
    Node 1 — Interactive questionnaire.
    Uses interrupt() multiple times — LangGraph replays resolved interrupts
    on each re-run, so only the next unanswered question pauses execution.
    temperature=0 on questionnaire generation ensures the same questions are
    produced each time the node is re-run after a resume.
    """

    # ── INTERRUPT 1: initial subject / grade intake ───────────────────────
    initial_input = interrupt({
        "node":  "scope_architect",
        "step":  "initial_intake",
        "prompt": (
            "╔══════════════════════════════════════════════════════════╗\n"
            "║    MULTI-AGENT TEXTBOOK CREATION SYSTEM  —  NODE 1      ║\n"
            "║              Scope & Curriculum Architect                ║\n"
            "╚══════════════════════════════════════════════════════════╝\n\n"
            "Welcome!  I will guide you through building a complete textbook.\n\n"
            "To begin, please tell me:\n"
            "  • What SUBJECT is this textbook for?\n"
            "  • What is the TARGET GRADE or AGE LEVEL?\n\n"
            "Examples:\n"
            "  'Biology for Grade 10'\n"
            "  'Python programming for adult beginners'\n"
            "  'Mathematics for Grade 3'\n"
        ),
    })

    # Parse subject + grade (temperature=0 → deterministic on re-runs)
    parsed = _llm_json(
        "Extract the subject and grade level from the user input. "
        "Also estimate the typical reading age in years (integer). "
        "Return JSON: "
        '{"subject": str, "grade_level": str, "estimated_reading_age": int}',
        str(initial_input),
        temperature=0.0,
    )
    subject      = parsed["subject"]
    grade_level  = parsed["grade_level"]
    reading_age  = int(parsed.get("estimated_reading_age", 12))

    # Generate questionnaire  (temperature=0 → same questions on every re-run)
    questionnaire = _llm_json(
        _QGEN_SYSTEM,
        f"Subject: {subject}\nGrade Level: {grade_level}",
        max_tokens=3000,
        temperature=0.0,
    )
    questions = questionnaire["questions"]
    q_count   = len(questions)

    # ── INTERRUPT 2: announce questionnaire plan ───────────────────────────
    interrupt({
        "node":  "scope_architect",
        "step":  "questionnaire_intro",
        "prompt": (
            f"\n✅  Scope identified: [{subject}]  for  [{grade_level}]\n\n"
            f"📋  I will ask you {q_count} targeted questions.\n"
            f"    Reason: {questionnaire['reasoning']}\n\n"
            "Press [Enter] or type 'ready' to begin..."
        ),
    })

    # ── INTERRUPTS 3 … N+2: one per questionnaire question ────────────────
    answers: dict[str, dict] = {}
    for i, q in enumerate(questions):
        opts = "\n".join(
            f"  [{chr(65 + j)}]  {opt}"
            + (
                f"\n       → {q['implications'][j]}"
                if j < len(q.get("implications", []))
                else ""
            )
            for j, opt in enumerate(q["options"])
        )
        answer = interrupt({
            "node":             "scope_architect",
            "step":             f"question_{i + 1}",
            "question_number":  i + 1,
            "total_questions":  q_count,
            "prompt": (
                f"\n{'━' * 56}\n"
                f"  Question {i + 1} / {q_count}\n"
                f"{'━' * 56}\n\n"
                f"  {q['text']}\n\n"
                f"{opts}\n\n"
                "  Enter your choice (A / B / C / D) or type a custom answer:"
            ),
        })
        answers[q["id"]] = {
            "question": q["text"],
            "answer":   str(answer).strip(),
            "category": q.get("category", "general"),
        }

    # ── INTERRUPT N+3: export format preference ────────────────────────────
    fmt_choice = interrupt({
        "node":  "scope_architect",
        "step":  "export_format",
        "prompt": (
            "\n📁  Final export format:\n"
            "  [A]  Microsoft Word  (.docx)   — editable, great for reviewers\n"
            "  [B]  PDF             (.pdf)    — fixed layout, print-ready\n\n"
            "  Enter A or B:"
        ),
    })
    export_fmt = "pdf" if str(fmt_choice).strip().upper() == "B" else "word"

    # ── Generate scope profile ─────────────────────────────────────────────
    scope_profile = _llm(
        _SCOPE_SYSTEM,
        (
            f"Subject: {subject}\n"
            f"Grade Level: {grade_level}\n"
            f"Estimated Reading Age: {reading_age}\n\n"
            f"Questionnaire Answers:\n{json.dumps(answers, indent=2)}"
        ),
        max_tokens=3000,
        temperature=0.7,
    )

    pedagogy  = next(
        (v["answer"] for v in answers.values() if v["category"] == "pedagogy"),
        "Direct Instruction",
    )
    standards = [
        v["answer"] for v in answers.values() if v["category"] == "standards"
    ]

    return {
        "subject":              subject,
        "grade_level":          grade_level,
        "target_reading_age":   reading_age,
        "pedagogical_style":    pedagogy,
        "compliance_standards": standards,
        "scope_profile":        scope_profile,
        "export_format":        export_fmt,
        "current_chapter_index": 0,
        "current_step":         "scaffolding",
        "approved_chapters":    [],   # initialise accumulator
        "research_cache":       {},
        "active_chapter_draft": None,
        "user_feedback_buffer": None,
        "selected_activity_type": None,
        "export_path":          None,
        "table_of_contents":    [],
        "total_chapters":       0,
    }


# ═════════════════════════════════════════════════════════════════════════════
# NODE 3 — STRUCTURAL SCAFFOLD  (Fully Automated)
# ═════════════════════════════════════════════════════════════════════════════

_SCAFFOLD_SYSTEM = """
You are a curriculum architect.  Build a complete, logically scaffolded
Table of Contents from the Project Scope Profile provided.

Rules:
• Each chapter must build on the prior chapter's concepts.
• Progress from FOUNDATIONAL → INTERMEDIATE → ADVANCED.
• Chapter count must match the scope profile's recommendation (8-15).
• Every chapter must have 3-4 subsection titles.

Return a JSON ARRAY (no other text) matching this exact per-chapter schema:
[
  {
    "chapter_id": 1,
    "title": "<Engaging chapter title>",
    "focus": "<One-sentence primary conceptual focus>",
    "prerequisites": ["<Concept A>", "<Concept B>"],
    "outcomes": [
      "Students will be able to ...",
      "Students will understand ...",
      "Students will apply ..."
    ],
    "subsections": [
      "<1.1 subsection title>",
      "<1.2 subsection title>",
      "<1.3 subsection title>"
    ],
    "status": "pending"
  }
]
"""


def structural_scaffold(state: TextbookSystemState) -> dict:
    """
    Node 3 — Fully automated TOC generator.
    No user interaction — reads scope_profile and produces the chapter array.
    """
    toc = _llm_json(
        _SCAFFOLD_SYSTEM,
        f"Project Scope Profile:\n{state['scope_profile']}",
        max_tokens=4000,
    )

    return {
        "table_of_contents": toc,
        "total_chapters":    len(toc),
        "current_step":      "gate",
    }


# ═════════════════════════════════════════════════════════════════════════════
# NODE 4a — CONTENT WRITER GATE  (Micro-Approval HITL)
# ═════════════════════════════════════════════════════════════════════════════

_STYLE_SAMPLE_SYSTEM = """
You are a textbook author.  For the chapter described, generate TWO contrasting
3-sentence prose style samples plus two activity module options.

Return JSON with EXACTLY this structure:
{
  "sample_A": {
    "style_label": "<e.g. Narrative / Storytelling>",
    "prose":       "<3-sentence sample in this style>",
    "activity":    "<e.g. Project-based learning>"
  },
  "sample_B": {
    "style_label": "<e.g. Academic / Socratic>",
    "prose":       "<3-sentence sample in this style>",
    "activity":    "<e.g. Conceptual essay questions>"
  }
}
"""


def writer_gate(state: TextbookSystemState) -> dict:
    """
    Node 4a — Micro-approval gate.
    Presents two style+activity options; user picks one before drafting begins.
    If this is a revision run, the previous feedback is displayed prominently.
    """
    idx     = state["current_chapter_index"]
    chapter = state["table_of_contents"][idx]

    # Generate two style samples (temperature=0 → stable on re-runs)
    samples = _llm_json(
        _STYLE_SAMPLE_SYSTEM,
        (
            f"Chapter title: {chapter['title']}\n"
            f"Focus: {chapter['focus']}\n"
            f"Grade level: {state['grade_level']}\n"
            f"Preferred pedagogy: {state['pedagogical_style']}"
        ),
        temperature=0.0,
    )

    is_revision = (
        state.get("current_step") == "revise"
        and state.get("user_feedback_buffer")
    )
    revision_banner = ""
    if is_revision:
        revision_banner = (
            f"\n⚠️   REVISION — The previous draft was rejected.\n"
            f"    Feedback: {state['user_feedback_buffer']}\n"
            f"    Adjust your style selection if needed.\n"
        )

    choice = interrupt({
        "node":       "writer_gate",
        "step":       "style_selection",
        "chapter_id": chapter["chapter_id"],
        "prompt": (
            f"\n╔══════════════════════════════════════════════════════════╗\n"
            f"║  NODE 4 · Writer Gate  │  Ch {chapter['chapter_id']} of {state['total_chapters']}                    ║\n"
            f"╚══════════════════════════════════════════════════════════╝\n"
            f"{revision_banner}\n"
            f"  📖  Chapter: {chapter['title']}\n"
            f"       Focus: {chapter['focus']}\n\n"
            f"Choose a writing style for this chapter:\n\n"
            f"  [A]  {samples['sample_A']['style_label']}\n"
            f"       \"{samples['sample_A']['prose']}\"\n"
            f"       Activity module: {samples['sample_A']['activity']}\n\n"
            f"  [B]  {samples['sample_B']['style_label']}\n"
            f"       \"{samples['sample_B']['prose']}\"\n"
            f"       Activity module: {samples['sample_B']['activity']}\n\n"
            "  Enter A or B:"
        ),
    })

    pick = "A" if str(choice).strip().upper() != "B" else "B"
    selected = samples[f"sample_{pick}"]

    return {
        "selected_activity_type": selected["activity"],
        "current_step":           "researching",
        "user_feedback_buffer":   None,   # clear after gate processes it
    }


# ═════════════════════════════════════════════════════════════════════════════
# NODE 2 — TARGETED RESEARCHER  (Fully Automated RAG Utility)
# ═════════════════════════════════════════════════════════════════════════════

_RESEARCH_SYSTEM = """
You are an academic research assistant performing a structured RAG-style
content retrieval for a textbook chapter.

Return a JSON object with EXACTLY this schema (no extra keys):
{
  "chapter_key": "<chapter_N>",
  "key_concepts": [
    {"term": str, "definition": str, "source_type": "<textbook|journal|standard>"}
  ],
  "real_world_analogies": [
    {"analogy": str, "age_appropriateness": str, "subject_link": str}
  ],
  "common_misconceptions": [
    {"misconception": str, "correction": str}
  ],
  "supporting_data": [
    {"fact": str, "plausible_source": str}
  ],
  "vocabulary_tier2": ["<academic word>"],
  "vocabulary_tier3": ["<domain-specific term>"],
  "cross_disciplinary_links": ["<connection to another field>"]
}

Match all analogies and vocabulary precisely to the grade level / reading age.
"""


def targeted_researcher(state: TextbookSystemState) -> dict:
    """
    Node 2 — Automated research utility.
    Populates the research_cache for the current chapter.
    Skips re-research if the cache entry already exists (e.g., on a revision
    loop) to avoid unnecessary API calls.
    """
    idx         = state["current_chapter_index"]
    chapter     = state["table_of_contents"][idx]
    chapter_key = f"chapter_{chapter['chapter_id']}"

    cache = dict(state.get("research_cache", {}))

    if chapter_key not in cache:
        data = _llm_json(
            _RESEARCH_SYSTEM,
            (
                f"Chapter: {chapter['title']}\n"
                f"Focus: {chapter['focus']}\n"
                f"Outcomes: {', '.join(chapter['outcomes'])}\n"
                f"Subsections: {', '.join(chapter['subsections'])}\n"
                f"Subject: {state['subject']}\n"
                f"Grade Level: {state['grade_level']}\n"
                f"Reading Age: {state['target_reading_age']}\n"
                f"Activity type: {state.get('selected_activity_type', 'review questions')}"
            ),
            max_tokens=3000,
        )
        cache[chapter_key] = data

    return {
        "research_cache": cache,
        "current_step":   "drafting",
    }


# ═════════════════════════════════════════════════════════════════════════════
# NODE 4b — CONTENT WRITER DRAFT  (Fully Automated)
# ═════════════════════════════════════════════════════════════════════════════

_DRAFT_SYSTEM = """
You are an expert textbook author.  Write a COMPLETE, publication-ready
chapter in Markdown format.

MANDATORY STRUCTURE (all sections required, in this order):

## Introduction
Hook + list of 3-5 learning objectives as bullet points.

## [Subsection 1 title from blueprint]
Full narrative prose.  Include real-world examples drawn from the research cache.

## [Subsection 2 title from blueprint]
...

## [Subsection 3 title from blueprint]
...

## Real-World Connections
2-3 age-appropriate analogies (box-style: **Analogy N:** text).

## Common Misconceptions
Table or bullet list: **Misconception:** text → **Correction:** text.

## Chapter Summary
Bulleted key takeaways (7-10 items).

## Review Questions
5 questions, labelled Q1–Q5, escalating in difficulty
(recall → comprehension → application → analysis → synthesis).

## Activity: [activity type]
Full instructions for the selected activity module.

## Vocabulary Builder
Definition list: **Term** — definition.

QUALITY CONSTRAINTS:
• Match the target reading age vocabulary precisely.
• Active voice throughout.
• Smooth transitions between every section.
• Do NOT include a chapter number heading — the exporter adds that.
• Bold important terms on first use.
"""


def writer_draft(state: TextbookSystemState) -> dict:
    """
    Node 4b — Full chapter draft.
    Synthesises research cache + chapter blueprint into publication-ready Markdown.
    On revision runs, the user feedback is appended as an explicit instruction.
    """
    idx         = state["current_chapter_index"]
    chapter     = state["table_of_contents"][idx]
    chapter_key = f"chapter_{chapter['chapter_id']}"
    research    = state["research_cache"].get(chapter_key, {})

    revision_note = ""
    if state.get("user_feedback_buffer"):
        revision_note = (
            f"\n\n━━━ REVISION INSTRUCTIONS ━━━\n"
            f"{state['user_feedback_buffer']}\n"
            f"Address every point above in this rewrite.\n"
        )

    draft = _llm(
        _DRAFT_SYSTEM,
        (
            f"CHAPTER BLUEPRINT\n"
            f"Title:         {chapter['title']}\n"
            f"Focus:         {chapter['focus']}\n"
            f"Prerequisites: {', '.join(chapter['prerequisites'])}\n"
            f"Outcomes:\n"
            + "\n".join(f"  • {o}" for o in chapter["outcomes"])
            + f"\nSubsections:   {', '.join(chapter['subsections'])}\n\n"
            f"SCOPE PARAMETERS\n"
            f"Subject:     {state['subject']}\n"
            f"Grade:       {state['grade_level']}\n"
            f"Reading Age: {state['target_reading_age']}\n"
            f"Pedagogy:    {state['pedagogical_style']}\n"
            f"Activity:    {state.get('selected_activity_type', 'Conceptual review questions')}\n\n"
            f"RESEARCH CACHE\n{json.dumps(research, indent=2)}"
            f"{revision_note}"
        ),
        max_tokens=6000,
        temperature=0.75,
    )

    return {
        "active_chapter_draft": draft,
        "current_step":         "reviewing",
    }


# ═════════════════════════════════════════════════════════════════════════════
# NODE 5 — QA & REVIEW EDITOR  (Binary HITL Gatekeeper)
# ═════════════════════════════════════════════════════════════════════════════

_QA_EVAL_SYSTEM = """
You are a senior academic editor.  Evaluate the textbook chapter draft and
return a JSON evaluation report with EXACTLY these keys:
{
  "flesch_kincaid_estimate": "<e.g. Grade 9.2>",
  "reading_level_verdict": "<PASS|FAIL — 1-sentence reason>",
  "structural_completeness": "<PASS|FAIL — list any missing mandatory sections>",
  "pedagogical_alignment": "<PASS|FAIL — matches specified style?>",
  "factual_integrity": "<PASS|WARN|FAIL — observations>",
  "auto_fixes_applied": ["<description of fix 1>", "..."],
  "editor_summary": "<3-4 sentence overall assessment>",
  "ai_recommendation": "<APPROVE|REVISE — with specific revision points if REVISE>"
}
"""

_QA_FIX_SYSTEM = """
You are a copy editor.  Apply the following corrections to the chapter draft
and return the COMPLETE corrected chapter text in Markdown.
Do not add or remove sections — only fix the issues listed.
Preserve all headings, lists, and formatting.
"""


def review_editor(state: TextbookSystemState) -> dict:
    """
    Node 5 — QA scoring + binary HITL gate.
    Step 1: evaluate and auto-fix the draft.
    Step 2: interrupt() → user sees QA report + chapter preview → APPROVE / REVISE.
    Step 3: if REVISE → second interrupt() to collect specific feedback.
    """
    idx     = state["current_chapter_index"]
    chapter = state["table_of_contents"][idx]
    draft   = state.get("active_chapter_draft") or ""

    # ── Step 1: QA evaluation ─────────────────────────────────────────────
    qa_eval = _llm_json(
        _QA_EVAL_SYSTEM,
        (
            f"SCOPE EXCERPT:\n{state['scope_profile'][:800]}\n\n"
            f"CHAPTER BLUEPRINT:\n"
            f"Title: {chapter['title']}\n"
            f"Grade: {state['grade_level']}\n"
            f"Pedagogy: {state['pedagogical_style']}\n\n"
            f"DRAFT CHAPTER:\n{draft}"
        ),
        max_tokens=2000,
    )

    # ── Step 2: Auto-fix pass ─────────────────────────────────────────────
    fixes_description = "; ".join(qa_eval.get("auto_fixes_applied", []))
    if fixes_description:
        corrected = _llm(
            _QA_FIX_SYSTEM,
            f"Fixes to apply: {fixes_description}\n\nOriginal chapter:\n{draft}",
            max_tokens=7000,
            temperature=0.2,
        )
    else:
        corrected = draft

    # ── Build QA report banner ────────────────────────────────────────────
    report = (
        f"\n╔══════════════════════════════════════════════════════════╗\n"
        f"║   NODE 5 · QA Review Editor  │  Ch {chapter['chapter_id']} of {state['total_chapters']}                ║\n"
        f"╚══════════════════════════════════════════════════════════╝\n\n"
        f"  📊  QUALITY CONTROL REPORT\n"
        f"  {'─' * 52}\n"
        f"  Reading Level:       {qa_eval.get('flesch_kincaid_estimate','—')}\n"
        f"  Level Compliance:    {qa_eval.get('reading_level_verdict','—')}\n"
        f"  Structure:           {qa_eval.get('structural_completeness','—')}\n"
        f"  Pedagogy Alignment:  {qa_eval.get('pedagogical_alignment','—')}\n"
        f"  Factual Integrity:   {qa_eval.get('factual_integrity','—')}\n"
        f"  Auto-fixes applied:  {fixes_description or 'None'}\n\n"
        f"  📝  Editor Summary:\n"
        f"  {qa_eval.get('editor_summary','').replace(chr(10), chr(10)+'  ')}\n\n"
        f"  🤖  AI Recommendation: {qa_eval.get('ai_recommendation','—')}\n"
    )

    preview_len = 1800
    preview = corrected[:preview_len]
    if len(corrected) > preview_len:
        preview += "\n  … [truncated — full draft committed on APPROVE] …"

    # ── INTERRUPT 1: present report + binary gate ─────────────────────────
    decision = interrupt({
        "node":       "review_editor",
        "step":       "binary_gate",
        "chapter_id": chapter["chapter_id"],
        "prompt": (
            report
            + f"\n  {'═' * 52}\n"
            f"  📄  CHAPTER PREVIEW: {chapter['title']}\n"
            f"  {'─' * 52}\n"
            f"{preview}\n"
            f"  {'─' * 52}\n\n"
            f"  DECISION REQUIRED:\n"
            f"    [APPROVE]  Accept this chapter and advance.\n"
            f"    [REVISE]   Send back for a rewrite.\n\n"
            "  Enter APPROVE or REVISE:"
        ),
    })

    if str(decision).strip().upper().startswith("A"):
        # ── APPROVE ───────────────────────────────────────────────────────
        return {
            "approved_chapters":     [corrected],  # operator.add appends
            "active_chapter_draft":  corrected,
            "current_chapter_index": idx + 1,      # advance to next chapter
            "current_step":          "approved",
            "user_feedback_buffer":  None,
        }

    # ── REVISE: collect feedback ──────────────────────────────────────────
    feedback = interrupt({
        "node":  "review_editor",
        "step":  "collect_feedback",
        "prompt": (
            "\n  📝  Please describe the required revisions:\n"
            "  (Be specific — the writing agent reads this directly.)\n\n"
            "  Your feedback:"
        ),
    })

    return {
        "user_feedback_buffer": str(feedback).strip(),
        "current_step":         "revise",
        "active_chapter_draft": corrected,
    }


# ═════════════════════════════════════════════════════════════════════════════
# NODE 6 — DOCUMENT EXPORT SPECIALIST  (Fully Automated)
# ═════════════════════════════════════════════════════════════════════════════

def doc_exporter(state: TextbookSystemState) -> dict:
    """
    Node 6 — Automated file compiler.
    Triggered only when current_chapter_index >= total_chapters.
    Generates DOCX or PDF depending on export_format flag.
    """
    chapters    = state.get("approved_chapters", [])
    toc         = state["table_of_contents"]
    fmt         = state.get("export_format", "word")

    timestamp   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_subj   = re.sub(r"[^a-zA-Z0-9]+", "_", state["subject"]).lower()
    filename    = f"textbook_{safe_subj}_{timestamp}"

    output_dir  = os.path.join(os.path.dirname(__file__), "output")
    os.makedirs(output_dir, exist_ok=True)

    if fmt == "pdf":
        path = _export_pdf(state, chapters, toc, output_dir, filename)
    else:
        path = _export_docx(state, chapters, toc, output_dir, filename)

    return {
        "export_path":  path,
        "current_step": "complete",
    }


# ─── DOCX export ─────────────────────────────────────────────────────────────

def _export_docx(state: dict, chapters: list, toc: list,
                 output_dir: str, filename: str) -> str:
    from docx import Document  # python-docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

    # ── Title page ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(state["subject"].title())
    run.bold = True
    run.font.size = Pt(28)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"A Comprehensive Textbook for {state['grade_level']}")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(
        f"Pedagogical Approach: {state['pedagogical_style']}\n"
        f"Grade Level: {state['grade_level']}\n"
        f"Generated: {datetime.datetime.now().strftime('%B %Y')}"
    ).font.size = Pt(11)
    doc.add_page_break()

    # ── Table of Contents page ────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    for ch in toc:
        p = doc.add_paragraph(
            f"Chapter {ch['chapter_id']}: {ch['title']}",
            style="List Number",
        )
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()

    # ── Chapter pages ─────────────────────────────────────────────────────
    for ch_meta, ch_text in zip(toc, chapters):
        doc.add_heading(
            f"Chapter {ch_meta['chapter_id']}: {ch_meta['title']}", level=1
        )
        # Learning outcomes block
        bp = doc.add_paragraph("Learning Outcomes")
        bp.runs[0].bold = True
        for out in ch_meta.get("outcomes", []):
            doc.add_paragraph(out, style="List Bullet")
        doc.add_paragraph()
        _md_to_docx(doc, ch_text)
        doc.add_page_break()

    path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(path)
    return path


def _md_to_docx(doc, markdown_text: str) -> None:
    """Convert a subset of Markdown to python-docx elements."""
    from docx.shared import Pt

    for line in markdown_text.splitlines():
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line).strip(),
                              style="List Number")
        elif line.strip() == "":
            doc.add_paragraph()
        else:
            # Regular paragraph — handle inline **bold**
            p = doc.add_paragraph()
            for j, part in enumerate(re.split(r"\*\*(.+?)\*\*", line)):
                run = p.add_run(part)
                run.bold = bool(j % 2)


# ─── PDF export ──────────────────────────────────────────────────────────────

def _export_pdf(state: dict, chapters: list, toc: list,
                output_dir: str, filename: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

    path = os.path.join(output_dir, f"{filename}.pdf")

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )

    SS = getSampleStyleSheet()
    title_s    = ParagraphStyle("TB_Title",  parent=SS["Title"],
                                fontSize=28, spaceAfter=12, alignment=TA_CENTER)
    sub_s      = ParagraphStyle("TB_Sub",    parent=SS["Normal"],
                                fontSize=14, spaceAfter=6,  alignment=TA_CENTER,
                                textColor=colors.HexColor("#444444"))
    h1_s       = ParagraphStyle("TB_H1",     parent=SS["Heading1"],
                                fontSize=18, spaceBefore=20, spaceAfter=10)
    h2_s       = ParagraphStyle("TB_H2",     parent=SS["Heading2"],
                                fontSize=14, spaceBefore=14, spaceAfter=8)
    h3_s       = ParagraphStyle("TB_H3",     parent=SS["Heading3"],
                                fontSize=12, spaceBefore=10, spaceAfter=6)
    body_s     = ParagraphStyle("TB_Body",   parent=SS["Normal"],
                                fontSize=11, spaceAfter=8, leading=16,
                                alignment=TA_JUSTIFY)
    bullet_s   = ParagraphStyle("TB_Bullet", parent=body_s,
                                leftIndent=14, bulletIndent=4)

    story = []

    # Title page
    story += [
        Spacer(1, 50*mm),
        Paragraph(state["subject"].title(), title_s),
        Spacer(1, 8*mm),
        Paragraph(f"A Comprehensive Textbook for {state['grade_level']}", sub_s),
        Spacer(1, 4*mm),
        Paragraph(
            f"Pedagogical Approach: {state['pedagogical_style']}  •  "
            f"Generated: {datetime.datetime.now().strftime('%B %Y')}",
            sub_s,
        ),
        PageBreak(),
    ]

    # TOC
    story.append(Paragraph("Table of Contents", h1_s))
    for ch in toc:
        story.append(
            Paragraph(f"Chapter {ch['chapter_id']}: {ch['title']}", body_s)
        )
    story.append(PageBreak())

    # Chapters
    for ch_meta, ch_text in zip(toc, chapters):
        story.append(
            Paragraph(
                f"Chapter {ch_meta['chapter_id']}: {ch_meta['title']}", h1_s
            )
        )
        story.append(Paragraph("<b>Learning Outcomes</b>", body_s))
        for out in ch_meta.get("outcomes", []):
            story.append(Paragraph(f"• {out}", bullet_s))
        story.append(Spacer(1, 4*mm))

        for line in ch_text.splitlines():
            if line.startswith("## "):
                story.append(Paragraph(line[3:].strip(), h2_s))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:].strip(), h3_s))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"• {line[2:].strip()}", bullet_s))
            elif re.match(r"^\d+\.\s", line):
                story.append(
                    Paragraph(re.sub(r"^\d+\.\s", "", line).strip(), bullet_s)
                )
            elif line.strip():
                formatted = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
                story.append(Paragraph(formatted, body_s))
            else:
                story.append(Spacer(1, 3*mm))

        story.append(PageBreak())

    doc.build(story)
    return path
